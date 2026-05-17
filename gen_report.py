# -*- coding: utf-8 -*-
"""Генератор отчёта по практике для проекта RobotPathCraft."""

import sys
import io

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# --- Параметры страницы ---
section = doc.sections[0]
section.page_width    = Cm(21.0)
section.page_height   = Cm(29.7)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(1.5)
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)

FONT   = 'Times New Roman'
SIZE   = Pt(14)
INDENT = Cm(1.25)


def _fix_font(run, bold=False, italic=False, size=None, font=FONT):
    run.font.name = font
    run.font.size = size or SIZE
    run.bold      = bold
    run.italic    = italic
    rPr    = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)
    rFonts.set(qn('w:cs'),    font)
    existing = rPr.find(qn('w:rFonts'))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)


def _para_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
              first_line=True, space_after=0):
    fmt = p.paragraph_format
    fmt.alignment           = align
    fmt.first_line_indent   = INDENT if first_line else None
    fmt.space_before        = Pt(0)
    fmt.space_after         = Pt(space_after)
    fmt.line_spacing_rule   = WD_LINE_SPACING.ONE_POINT_FIVE


def add_heading(text):
    p   = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.alignment         = WD_ALIGN_PARAGRAPH.CENTER
    fmt.first_line_indent = None
    fmt.space_before      = Pt(12)
    fmt.space_after       = Pt(6)
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    _fix_font(run, bold=True)


def add_body(text, first_line=True):
    p = doc.add_paragraph()
    _para_fmt(p, first_line=first_line)
    run = p.add_run(text)
    _fix_font(run)


def add_bullet(text):
    p   = doc.add_paragraph(style='List Bullet')
    fmt = p.paragraph_format
    fmt.alignment         = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt.space_before      = Pt(0)
    fmt.space_after       = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.left_indent       = Cm(1.25)
    fmt.first_line_indent = None
    run = p.add_run(text)
    _fix_font(run)


def add_code(lines):
    for line in lines:
        p   = doc.add_paragraph()
        fmt = p.paragraph_format
        fmt.alignment         = WD_ALIGN_PARAGRAPH.LEFT
        fmt.first_line_indent = None
        fmt.left_indent       = Cm(1.25)
        fmt.space_before      = Pt(0)
        fmt.space_after       = Pt(0)
        fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(11)


def add_caption(text):
    p   = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.alignment         = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt.first_line_indent = INDENT
    fmt.space_before      = Pt(6)
    fmt.space_after       = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    _fix_font(run, italic=True)


def add_empty():
    p   = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before      = Pt(0)
    fmt.space_after       = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    _fix_font(p.add_run(''))


# ================================================================
# Титул
# ================================================================
p   = doc.add_paragraph()
fmt = p.paragraph_format
fmt.alignment         = WD_ALIGN_PARAGRAPH.CENTER
fmt.first_line_indent = None
fmt.space_before      = Pt(0)
fmt.space_after       = Pt(6)
fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
run = p.add_run('Отчёт по практике')
_fix_font(run, bold=True, size=Pt(16))

p2   = doc.add_paragraph()
fmt2 = p2.paragraph_format
fmt2.alignment         = WD_ALIGN_PARAGRAPH.CENTER
fmt2.first_line_indent = None
fmt2.space_before      = Pt(0)
fmt2.space_after       = Pt(18)
fmt2.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
run2 = p2.add_run(
    'Тема: разработка десктопного приложения '
    'для планирования траекторий мехатронных систем '
    '(RobotPathCraft)'
)
_fix_font(run2)

# ================================================================
# 1. Постановка задачи
# ================================================================
add_heading('1. Постановка задачи на практику')

add_body(
    'Планирование бесстолкновительных траекторий движения мобильных роботов '
    'и мехатронных систем в средах с препятствиями является одной из '
    'центральных задач современной робототехники. Наличие эффективного '
    'инструмента для автоматического построения оптимальных маршрутов '
    'позволяет сократить время разработки, повысить безопасность автономных '
    'систем и упростить прототипирование навигационных алгоритмов.'
)
add_body(
    'В рамках данной практики поставлена задача разработки программного '
    'продукта, реализующего интерактивное планирование траекторий в '
    'двумерном рабочем пространстве с препятствиями. Программа должна '
    'поддерживать несколько форматов карт, два алгоритма поиска пути, '
    'сглаживание найденного маршрута и предоставлять аналитические '
    'рекомендации оператору.'
)
add_body('Основные понятия предметной области:')
add_empty()

concepts = [
    ('Карта занятости (occupancy grid)',
     ' — дискретное представление рабочего пространства в виде двумерной '
     'матрицы, где каждая ячейка принимает значение 0 (свободна) или 1 '
     '(занята препятствием). Используется как входные данные для алгоритмов '
     'планирования пути.'),
    ('Алгоритм A*',
     ' — информированный алгоритм поиска кратчайшего пути на взвешенном '
     'графе, совмещающий реальную стоимость пройденного пути (g-функция) и '
     'эвристическую оценку расстояния до цели (h-функция). Гарантирует '
     'нахождение оптимального пути при допустимой эвристике.'),
    ('Алгоритм RRT (Rapidly-exploring Random Trees)',
     ' — вероятностный алгоритм планирования траектории, строящий дерево '
     'случайных выборок в пространстве конфигураций. Эффективен в средах со '
     'сложной геометрией препятствий и не требует дискретизации карты.'),
    ('Эвристическая функция',
     ' — приближённая оценка стоимости достижения цели из текущего узла, '
     'направляющая поиск в алгоритме A*. В данном проекте используется '
     'евклидово расстояние до целевой точки.'),
    ('Сглаживание пути (path smoothing)',
     ' — постобработка найденной траектории, устраняющая избыточные '
     'промежуточные точки при наличии прямой видимости между несмежными '
     'узлами пути. Позволяет сократить число поворотов и общую длину '
     'маршрута.'),
    ('Зазор безопасности',
     ' — минимальное расстояние от точек траектории до ближайшего '
     'препятствия. Является ключевым показателем безопасности движения: '
     'малое значение зазора указывает на необходимость снижения скорости '
     'или пересмотра маршрута.'),
    ('Пространство конфигураций',
     ' — абстрактное пространство, охватывающее все возможные положения и '
     'ориентации мобильного агента в рабочей среде. Служит теоретической '
     'основой для формализации задач планирования траекторий.'),
]

for term, definition in concepts:
    p   = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.alignment         = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt.first_line_indent = INDENT
    fmt.space_before      = Pt(0)
    fmt.space_after       = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r1 = p.add_run(term)
    _fix_font(r1, bold=True)
    r2 = p.add_run(definition)
    _fix_font(r2)

add_empty()

# ================================================================
# 2. Описание программного продукта
# ================================================================
add_heading('2. Описание программного продукта и используемых технологий')

add_body(
    'RobotPathCraft — десктопное приложение с графическим пользовательским '
    'интерфейсом (GUI), предназначенное для интерактивного планирования '
    'траекторий движения мобильных роботов в средах с препятствиями. '
    'Программа обеспечивает полный цикл работы: от загрузки карты рабочего '
    'пространства до получения аналитического отчёта о найденном маршруте.'
)
add_body('Основные задачи, которые решает программный продукт:')

tasks = [
    'загрузка карты рабочего пространства в одном из поддерживаемых форматов '
    '(PNG, JPG, JSON, CSV) и её преобразование в матрицу занятости;',
    'интерактивное задание начальной и конечной точек движения путём ввода '
    'координат или щелчка мышью по карте;',
    'планирование бесстолкновительного пути с использованием алгоритма A* '
    'или RRT;',
    'опциональное сглаживание траектории методом shortcut smoothing;',
    'визуализация карты, траектории и (для RRT) дерева случайных выборок;',
    'вычисление метрик пути: длины, числа поворотов, минимального зазора '
    'до препятствий и расчётного времени движения;',
    'формирование текстовых рекомендаций по безопасности и оптимальности '
    'маршрута;',
    'сохранение сеанса работы (карта и параметры) в формате JSON и экспорт '
    'изображения визуализации в PNG.',
]
for t in tasks:
    add_bullet(t)

add_empty()
add_body(
    'Форма программного продукта: десктопное GUI-приложение с разделённым '
    'интерфейсом, включающим панель управления параметрами слева и холст '
    'визуализации с полосой прогресса и панелью рекомендаций справа.'
)
add_empty()
add_body('Технологии и инструменты, использованные при разработке:')

tech = [
    ('Python 3.14',
     ' — основной язык программирования; выбран за богатую экосистему '
     'научных и GUI-библиотек.'),
    ('PyQt6 (≥ 6.4.0)',
     ' — фреймворк для построения кроссплатформенного GUI; обеспечивает '
     'главное окно, панели управления, диалоги и многопоточную обработку '
     'через QThread.'),
    ('NumPy (≥ 1.24.0)',
     ' — библиотека для работы с числовыми массивами; используется для '
     'хранения матрицы занятости и векторных вычислений при расчёте '
     'зазора безопасности.'),
    ('Matplotlib (≥ 3.7.0)',
     ' — библиотека визуализации; встроена в PyQt6-виджет (FigureCanvas) '
     'для отображения карты, траектории и дерева RRT.'),
    ('Pillow (≥ 9.5.0)',
     ' — библиотека обработки изображений; применяется для загрузки '
     'растровых карт (PNG, JPG) и их бинаризации по заданному порогу '
     'яркости.'),
    ('Стандартная библиотека Python',
     ' — модули heapq (приоритетная очередь для A*), math, dataclasses, '
     'json, pathlib.'),
]
for name, desc in tech:
    p   = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.alignment         = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt.first_line_indent = INDENT
    fmt.space_before      = Pt(0)
    fmt.space_after       = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r1 = p.add_run(name)
    _fix_font(r1, bold=True)
    r2 = p.add_run(desc)
    _fix_font(r2)

add_empty()

# ================================================================
# 3. Структура и функциональность
# ================================================================
add_heading('3. Структура и описание функциональности программного продукта')

add_body(
    'Программный продукт имеет модульную структуру, разделённую по '
    'функциональным областям. Корневой каталог содержит точку входа main.py '
    'и файл зависимостей requirements.txt. Функциональные подсистемы '
    'выделены в отдельные пакеты: gui, algorithms, analysis, map_loader.'
)
add_empty()
add_body('Состав модулей и их назначение:')

modules = [
    ('main.py',
     ' — точка входа; инициализирует QApplication, создаёт и показывает '
     'главное окно. При критической ошибке выводит диалог с трассировкой.'),
    ('gui/main_window.py',
     ' — главное окно приложения; координирует взаимодействие всех '
     'компонентов: загрузку карты, запуск планировщика в фоновом потоке '
     '(PlannerWorker : QObject), отображение статистики и рекомендаций.'),
    ('gui/canvas_widget.py',
     ' — виджет визуализации на основе FigureCanvas (Matplotlib); '
     'отображает карту занятости, маркеры старта и цели, найденный путь '
     'и дерево RRT.'),
    ('gui/control_panel.py',
     ' — левая панель управления; содержит элементы ввода координат, '
     'выбора алгоритма, настройки параметров движения и кнопки действий.'),
    ('algorithms/astar.py',
     ' — реализация алгоритма A* с восьмисвязной сеткой и евклидовой '
     'эвристикой; поддерживает настраиваемый шаг сетки и радиус достижения '
     'цели.'),
    ('algorithms/rrt.py',
     ' — реализация алгоритма RRT с эвристическим смещением выборки к цели '
     '(goal bias 15 %) и проверкой свободности прямолинейного отрезка.'),
    ('algorithms/smoother.py',
     ' — сглаживание пути методом shortcut smoothing: удаляет промежуточные '
     'точки при наличии прямой видимости между несмежными узлами.'),
    ('analysis/recommender.py',
     ' — вычисление метрик траектории (длина, число поворотов, минимальный '
     'зазор, расчётное время) и формирование текстовых рекомендаций.'),
    ('map_loader/',
     ' — пакет загрузки карт; включает три загрузчика (csv_loader, '
     'json_loader, image_loader) и контейнер данных MapData.'),
]

for name, desc in modules:
    p   = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.alignment         = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt.first_line_indent = INDENT
    fmt.space_before      = Pt(0)
    fmt.space_after       = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r1 = p.add_run(name)
    _fix_font(r1, bold=True)
    r2 = p.add_run(desc)
    _fix_font(r2)

add_empty()
add_body(
    'Ниже представлены три ключевых фрагмента программного кода с описанием '
    'алгоритмов работы.'
)
add_empty()

# --- Листинг 1: A* ---
add_caption('Листинг 1. Поиск пути алгоритмом A* (algorithms/astar.py)')
add_code([
    'def astar_search(grid, start, goal, goal_radius=0.0, step=1):',
    '    h, w = grid.shape',
    '    open_heap = []',
    '    heapq.heappush(open_heap, (0.0, start))',
    '    came_from = {}',
    '    g_score   = {start: 0.0}',
    '    visited   = set()',
    '    while open_heap:',
    '        _, current = heapq.heappop(open_heap)',
    '        if current in visited:',
    '            continue',
    '        visited.add(current)',
    '        if _heuristic(current, goal) <= goal_radius:',
    '            path = [current]',
    '            while current in came_from:',
    '                current = came_from[current]',
    '                path.append(current)',
    '            path.reverse()',
    '            return path',
    '        cx, cy = current',
    '        for nx, ny in _neighbors((cx, cy), h, w):',
    '            tentative = g_score[current] + math.hypot(nx-cx, ny-cy)',
    '            if tentative < g_score.get((nx, ny), float("inf")):',
    '                came_from[(nx, ny)] = current',
    '                g_score[(nx, ny)]   = tentative',
    '                f = tentative + _heuristic((nx, ny), goal)',
    '                heapq.heappush(open_heap, (f, (nx, ny)))',
    '    return None',
])
add_empty()
add_body(
    'Алгоритм A* обходит узлы сетки в порядке возрастания оценки '
    'f = g + h, где g — стоимость пройденного пути, h — евклидова '
    'эвристика. Словарь came_from хранит родителя каждого посещённого '
    'узла, что позволяет восстановить путь в обратном порядке при '
    'достижении цели. Алгоритм поддерживает восьмисвязную сетку, '
    'допуская диагональные переходы с весом sqrt(2).'
)
add_empty()

# --- Листинг 2: RRT ---
add_caption('Листинг 2. Ключевой цикл алгоритма RRT (algorithms/rrt.py)')
add_code([
    'for _ in range(max_iter):',
    '    # С вероятностью 15% направляем выборку к цели (goal bias)',
    '    if random.random() < 0.15:',
    '        sample = (float(goal[0]), float(goal[1]))',
    '    else:',
    '        sample = (random.uniform(0, w-1), random.uniform(0, h-1))',
    '    nearest_idx = min(range(len(nodes)),',
    '        key=lambda i: _distance((nodes[i].x, nodes[i].y), sample))',
    '    nearest = nodes[nearest_idx]',
    '    angle   = math.atan2(sample[1]-nearest.y, sample[0]-nearest.x)',
    '    new_x   = nearest.x + step_size * math.cos(angle)',
    '    new_y   = nearest.y + step_size * math.sin(angle)',
    '    if not _is_free(grid, new_x, new_y):',
    '        continue',
    '    if not _line_collision_free(grid, (nearest.x, nearest.y),',
    '                                      (new_x, new_y)):',
    '        continue',
    '    nodes.append(Node(new_x, new_y, nearest_idx))',
    '    if _distance((new_x, new_y), (goal[0], goal[1])) <= goal_radius:',
    '        # Путь найден — восстанавливаем обход по дереву',
    '        ...',
])
add_empty()
add_body(
    'RRT строит дерево случайных выборок, постепенно заполняя пространство '
    'конфигураций. На каждой итерации генерируется случайная точка (с '
    'вероятностью 15 % вместо неё используется координата цели — goal bias). '
    'Дерево расширяется в направлении этой точки на шаг step_size с '
    'обязательной проверкой свободности отрезка. При входе в радиус '
    'достижения цели путь восстанавливается обходом дерева по ссылкам на '
    'родительские узлы.'
)
add_empty()

# --- Листинг 3: recommender ---
add_caption(
    'Листинг 3. Формирование рекомендаций по траектории '
    '(analysis/recommender.py)'
)
add_code([
    'def build_recommendation(grid, path, speed, min_turn_radius):',
    '    length      = path_length(path)',
    '    turns       = turn_count(path)',
    '    clearance   = min_clearance(grid, path)',
    '    travel_time = length / speed if speed > 0 else float("inf")',
    '    lines = [',
    '        f"Длина пути:           {length:.2f} ед.",',
    '        f"Расчётное время:      {travel_time:.2f} с",',
    '        f"Количество поворотов: {turns}",',
    '        f"Мин. зазор:           {clearance:.2f} ед.",',
    '    ]',
    '    if turns > 12:',
    '        lines.append("Много поворотов: рассмотрите сглаживание.")',
    '    if clearance < 2.0:',
    '        lines.append("Критически малый зазор: снизьте скорость.")',
    '    elif clearance < min_turn_radius:',
    '        lines.append("Узкие участки: рекомендуется осторожный проход.")',
    '    return "\\n".join(lines)',
])
add_empty()
add_body(
    'Модуль анализа вычисляет четыре ключевые метрики маршрута и на их '
    'основе формирует адресные текстовые рекомендации. Функция min_clearance '
    'использует векторные операции NumPy для вычисления евклидовых расстояний '
    'от каждой точки пути до всего множества препятствий, что позволяет '
    'своевременно предупредить оператора об опасных участках маршрута.'
)
add_empty()

# Заглушки для скриншотов
add_body(
    'Рисунок 1. Главное окно приложения RobotPathCraft с загруженной картой '
    'и найденным маршрутом алгоритма A* (вставить скриншот).'
)
add_empty()
add_body(
    'Рисунок 2. Результат работы алгоритма RRT: визуализация дерева '
    'случайных выборок и найденной траектории (вставить скриншот).'
)
add_empty()
add_body(
    'Рисунок 3. Панель аналитических рекомендаций с вычисленными метриками '
    'пути (вставить скриншот).'
)
add_empty()

# ================================================================
# 4. Заключение
# ================================================================
add_heading('4. Заключение')

add_body(
    'В рамках практики был разработан программный продукт RobotPathCraft — '
    'десктопное GUI-приложение для планирования траекторий движения '
    'мехатронных систем в средах с препятствиями. Продукт реализует два '
    'классических алгоритма планирования пути (A* и RRT), поддерживает '
    'загрузку карт в форматах PNG, JPG, JSON и CSV, а также обеспечивает '
    'постобработку траектории методом shortcut smoothing.'
)
add_body('В ходе работы были достигнуты следующие результаты:')

conclusions = [
    'реализован и проверен алгоритм A* с восьмисвязной сеткой и евклидовой '
    'эвристикой, обеспечивающий нахождение оптимального пути;',
    'реализован алгоритм RRT с эвристическим смещением к цели (goal bias), '
    'позволяющий строить траектории в средах со сложной геометрией препятствий;',
    'разработан модуль анализа, вычисляющий длину пути, число поворотов, '
    'минимальный зазор безопасности и расчётное время движения с '
    'формированием адресных рекомендаций оператору;',
    'создан GUI на базе PyQt6 с возможностью интерактивного выбора точек, '
    'масштабирования карты, сохранения сеансов и экспорта визуализации.',
]
for c in conclusions:
    add_bullet(c)

add_empty()
add_body(
    'Полученные результаты демонстрируют применимость классических '
    'алгоритмов планирования пути в интерактивных инструментах для '
    'инженерного проектирования. Разработанный продукт может использоваться '
    'как в учебных целях, так и в качестве базы для создания более сложных '
    'систем навигации с поддержкой кинематических ограничений робота и '
    'трёхмерных карт.'
)
add_body(
    'В процессе выполнения практики были закреплены знания в области '
    'алгоритмов планирования пути, разработки GUI-приложений на Python, '
    'работы с NumPy и Matplotlib, а также получены практические навыки '
    'проектирования модульных программных систем.'
)
add_empty()

# ================================================================
# 5. Список источников
# ================================================================
add_heading('5. Список использованных источников')

sources = [
    '1. LaValle S. M. Planning Algorithms. — Cambridge University Press, '
    '2006. — 842 p.',
    '2. Russell S., Norvig P. Artificial Intelligence: A Modern Approach. '
    '4th ed. — Pearson, 2020. — 1132 p.',
    '3. Cormen T. H. и др. Introduction to Algorithms. 4th ed. — '
    'MIT Press, 2022. — 1312 p.',
    '4. Karaman S., Frazzoli E. Sampling-based Algorithms for Optimal '
    'Motion Planning // The International Journal of Robotics Research. — '
    '2011. — Vol. 30, № 7. — P. 846–894.',
    '5. Документация PyQt6. — URL: '
    'https://www.riverbankcomputing.com/static/Docs/PyQt6 '
    '(дата обращения: 05.05.2026).',
    '6. Документация NumPy. — URL: https://numpy.org/doc/stable '
    '(дата обращения: 05.05.2026).',
    '7. Документация Matplotlib. — URL: https://matplotlib.org/stable '
    '(дата обращения: 05.05.2026).',
]
for src in sources:
    p   = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.alignment         = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt.first_line_indent = None
    fmt.left_indent       = Cm(1.25)
    fmt.space_before      = Pt(0)
    fmt.space_after       = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(src)
    _fix_font(run)

# ================================================================
out = r'C:\Users\Ilya\Desktop\Отчёт_RobotPathCraft.docx'
doc.save(out)
print('Saved:', out)
